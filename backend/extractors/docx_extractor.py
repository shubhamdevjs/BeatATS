"""DOCX text extraction using python-docx."""
from docx import Document
from docx.oxml.ns import qn
from typing import Dict, List, Any
from pathlib import Path


def extract_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract text and metadata from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dict containing:
            - text: Full extracted text
            - pages: Single page with all text (DOCX doesn't have pages)
            - page_count: Always 1 for DOCX
            - signals: Layout detection signals
            - hyperlinks: Extracted hyperlinks
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    doc = Document(file_path)
    
    paragraphs = []
    hyperlinks = []
    signals = {
        "likely_two_column": False,
        "contains_tables": len(doc.tables) > 0,
        "has_icons": False,
        "low_text_density": False
    }
    
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
            
        # Extract hyperlinks from paragraph
        for hyperlink in para._element.findall(qn('w:hyperlink')):
            link_text = "".join(node.text for node in hyperlink.iter() if node.text)
            # Get the relationship ID
            r_id = hyperlink.get(qn('r:id'))
            if r_id and link_text:
                try:
                    url = doc.part.rels[r_id].target_ref
                    hyperlinks.append({
                        "text": link_text,
                        "url": url
                    })
                except (KeyError, AttributeError):
                    pass
    
    # Extract text from tables
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(" | ".join(row_text))
    
    full_text = "\n".join(paragraphs)
    if table_text:
        full_text += "\n\n" + "\n".join(table_text)
    
    # Check text density
    if len(full_text.strip()) < 200:
        signals["low_text_density"] = True
    
    return {
        "text": full_text,
        "pages": [{"page": 1, "text": full_text}],
        "page_count": 1,
        "signals": signals,
        "hyperlinks": hyperlinks
    }
